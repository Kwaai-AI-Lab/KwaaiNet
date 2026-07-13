"""
Rebuild RAGPerformanceReport-20260712.docx from _base.docx (fresh pandoc output)
so images never accumulate across runs.

Column width units: twips (1 inch = 1440 twips), written to both
w:tblGrid/w:gridCol and w:tcPr/w:tcW so Word respects them.

User-calibrated widths from their T0–T2 edits:
  numeric cols (Chunks/Entities/Speed/%): 0.62–0.88in
  KB names: 1.2in
  Notes/long-text: remaining (~4–6in)
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

TWIPS = 1440   # twips per inch

def set_col_widths(table, widths_in):
    """Write column widths (inches) to tblGrid and every tcW cell."""
    tw = [int(w * TWIPS) for w in widths_in]

    # 1. Update tblGrid
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tblGrid)
    gridCols = tblGrid.findall(qn("w:gridCol"))
    # Rebuild grid cols to match our count
    for gc in gridCols:
        tblGrid.remove(gc)
    for w in tw:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    # 2. Update every cell's tcW
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if i >= len(tw):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(tw[i]))
            tcW.set(qn("w:type"), "dxa")

def set_table_font(table, size_pt):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)

def insert_image_after(doc, elem, img_path, width_in, caption=None):
    body = doc.element.body
    idx  = list(body).index(elem)

    tmp  = doc.add_paragraph()
    run  = tmp.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    img_p = tmp._p
    img_p.getparent().remove(img_p)

    # Centre the image
    pPr = img_p.get_or_add_pPr()
    jc  = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    pPr.append(jc)
    body.insert(idx + 1, img_p)

    if caption:
        cap = OxmlElement("w:p")
        pPr2 = OxmlElement("w:pPr")
        jc2  = OxmlElement("w:jc"); jc2.set(qn("w:val"), "center")
        pPr2.append(jc2); cap.append(pPr2)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = caption; r.append(t)
        cap.append(r)
        body.insert(idx + 2, cap)

def find_heading(doc, fragment):
    for p in doc.paragraphs:
        if fragment in p.text and p.style.name.startswith("Heading"):
            return p._p
    return None

def tbls(doc):
    return [b for b in doc.element.body if b.tag.split("}")[1] == "tbl"]

# ── Open fresh base (never the output file) ───────────────────────────────────
doc = Document("_base.docx")

# ── Landscape ─────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
# Letter landscape: 11 × 8.5 in expressed in EMU (python-docx unit)
sec.page_width  = Inches(11)
sec.page_height = Inches(8.5)
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(0.75)

# ── Text fix: status line ─────────────────────────────────────────────────────
for para in doc.paragraphs:
    if "code complete, not yet committed" in para.text:
        for run in para.runs:
            run.text = run.text.replace(
                "Status: code complete, not yet committed.",
                "Status: committed (d651483) and tagged v0.5.1."
            )

# ── D6 row in Table 2 (Pipeline KBs) ─────────────────────────────────────────
d6_row = ["D6 ★", "1,152", "911", "35 min", "0.54 c/s", "89.5%", "N/A¹"]
t2 = doc.tables[2]
row = t2.add_row()
for cell, val in zip(row.cells, d6_row):
    cell.text = val
    if val == "D6 ★":
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

# ── Column widths (twips via set_col_widths) ──────────────────────────────────
# Widths derived from user's T0–T2 edits:
#   numeric: 0.62–0.88in  |  names: 1.2–1.75in  |  notes: 4–6in

# T0  5c: Run / Machine / Score / Questions / Notes
set_col_widths(doc.tables[0], [1.01, 1.75, 0.81, 0.94, 4.64])

# T1  4c: Phase / Entity Types / Context Window / Recall
set_col_widths(doc.tables[1], [1.65, 4.30, 2.00, 1.20])

# T2  7c: KB / Chunks / Entities / Build / Speed / Retrieval / Generation
set_col_widths(doc.tables[2], [1.20, 0.62, 0.62, 0.81, 0.69, 0.88, 4.33])
set_table_font(doc.tables[2], 9.5)

# T3  8c: KB / Chunks / Entities / Build / Speed / Retrieval / Generation / Eval method
# Same numeric pattern as T2; Eval method takes the Notes role
set_col_widths(doc.tables[3], [1.20, 0.62, 0.62, 0.81, 0.69, 0.88, 0.88, 3.45])
set_table_font(doc.tables[3], 9.5)

# T4  4c: Machine / GPU / Status / Notes
set_col_widths(doc.tables[4], [1.75, 1.50, 0.81, 5.09])

# T5  3c: Domain type / Retrieval range / Notes
set_col_widths(doc.tables[5], [2.00, 1.00, 6.15])

# T6  3c: Entity type / Per-type recall / Notes
set_col_widths(doc.tables[6], [1.20, 1.00, 6.95])

# T7  3c: Item / Priority / Notes
set_col_widths(doc.tables[7], [2.00, 0.81, 6.34])

# ── Insert charts (once, from clean base) ─────────────────────────────────────
# Process bottom-up so earlier inserts don't shift indices of later ones.

tbl_list = tbls(doc)

# Fig 5 → after T6 (entity extraction per type)
insert_image_after(doc, tbl_list[6], "chart5_entity_density.png", 7.5,
                   "Figure 5 — Entity extraction density across pipeline KBs + D6")

# Fig 4 → after T1 (D6 phase progression)
insert_image_after(doc, tbl_list[1], "chart4_d6_phases.png", 6.5,
                   "Figure 4 — D6 retrieval recall by entity extraction phase")

# Refresh after two inserts
tbl_list = tbls(doc)

# Fig 1 → after T2 (pipeline KBs)
insert_image_after(doc, tbl_list[2], "chart1_recall_pipeline.png", 8.5,
                   "Figure 1 — Retrieval & generation recall, pipeline KBs + D6")

# Refresh
tbl_list = tbls(doc)

# Fig 2 → after T3 (pre-pipeline KBs)
insert_image_after(doc, tbl_list[3], "chart2_recall_all.png", 8.5,
                   "Figure 2 — Retrieval recall across all 14 knowledge bases (ranked)")

# Fig 3 → after last bullet under "Speed observations"
speed_h = find_heading(doc, "Speed observations")
if speed_h is not None:
    body_list = list(doc.element.body)
    idx = body_list.index(speed_h)
    last = idx
    for j in range(idx + 1, len(body_list)):
        tag = body_list[j].tag.split("}")[1]
        if tag == "p":
            ps = body_list[j].find(".//" + qn("w:pStyle"))
            if ps is not None and ps.get(qn("w:val"), "").startswith("Heading"):
                break
            last = j
    insert_image_after(doc, body_list[last], "chart3_build_time.png", 7.5,
                       "Figure 3 — Build time vs corpus size (bubble = entity count, colour = throughput)")

doc.save("RAGPerformanceReport-20260712.docx")
print("Saved.")
